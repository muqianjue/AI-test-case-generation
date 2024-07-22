import os
import tempfile
from urllib.parse import urlparse
import requests
import base64
import json
import logging
from io import BytesIO
from docx import Document
from docx.parts.image import ImagePart
from PIL import Image

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class WordExtractor:
    """Load DOCX files and perform OCR on images within the document."""

    def __init__(self, file_path: str, paddleocr_url: str = "http://10.12.3.94:8956/ocr"):
        """Initialize WordExtractor with file path and optional OCR URL."""
        self.file_path = file_path
        self.paddleocr_url = paddleocr_url

        if "~" in self.file_path:
            self.file_path = os.path.expanduser(self.file_path)

        # If the file is a web path, download it to a temporary file, and use that
        if not os.path.isfile(self.file_path) and self._is_valid_url(self.file_path):
            r = requests.get(self.file_path)

            if r.status_code != 200:
                raise ValueError(
                    f"Check the url of your file; returned status code {r.status_code}"
                )

            self.web_path = self.file_path
            self.temp_file = tempfile.NamedTemporaryFile(delete=False)
            self.temp_file.write(r.content)
            self.temp_file.close()
            self.file_path = self.temp_file.name

        elif not os.path.isfile(self.file_path):
            raise ValueError(
                f"File path {self.file_path} is not a valid file or url"
            )

    def __del__(self):
        """Cleanup method to delete temporary file if it exists."""
        if hasattr(self, "temp_file"):
            os.remove(self.temp_file.name)

    def extract(self) -> list:
        """Extract text, OCR processed images, and table data from the DOCX file."""
        doc = Document(self.file_path)
        text_content = []
        image_content = []
        table_content = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text.strip())

        # Extract OCR processed text from images
        for rel in doc.part.rels.values():
            if isinstance(rel.target_part, ImagePart):
                result = self._pilimg_to_text(Image.open(BytesIO(rel.target_part.blob)), self.paddleocr_url)
                if result:
                    image_content.append(result)

        # Extract table data
        for table in doc.tables:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            table_data = []
            if len(table.rows) == 1:
                table_data.append(",".join(headers))
            for row in table.rows[1:]:
                row_content = []
                for column_index, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    row_content.append(f'"{headers[column_index]}":"{cell_text}"')
                table_data.append("{" + ",".join(row_content) + "}")
            table_content.append(table_data)

        return {
            "text_content": text_content,
            "image_content": image_content,
            "table_content": table_content}

    def _is_valid_url(self, url: str) -> bool:
        """Check if the url is valid."""
        parsed = urlparse(url)
        return bool(parsed.netloc) and bool(parsed.scheme)

    def _pilimg_to_text(self, pil_img, paddleocr_url: str) -> str:
        """Convert PIL image to OCR processed text using PaddleOCR."""
        buffered = BytesIO()
        pil_img.save(buffered, format=pil_img.format)
        image_data_bytes = buffered.getvalue()
        img_base64 = base64.b64encode(image_data_bytes).decode('utf-8')
        return self._get_ocr_result(img_base64, paddleocr_url)

    def _get_ocr_result(self, img_base64: str, paddleocr_url: str) -> str:
        """Send image base64 data to PaddleOCR API and retrieve OCR results."""
        headers = {"Content-Type": "application/json"}
        data = {"img_base64": img_base64}
        try:
            response = requests.post(url=paddleocr_url, headers=headers, data=json.dumps(data))
            response.raise_for_status()
            result = response.json()["rec_texts"][0]
            return ",".join(result)
        except requests.exceptions.RequestException as e:
            logging.error("Error performing OCR: %s", str(e))
            return ""

# 示例使用
if __name__ == "__main__":
    # 替换为你的 DOCX 文件路径
    docx_file = r"D:\TestCases\航空云监控管理_需求规格说明书20240223.docx"

    # 创建 WordExtractor 实例
    extractor = WordExtractor(docx_file)

    try:
        # 提取文档内容
        extracted_content = extractor.extract()
        print("提取的文本内容和OCR处理结果：")
        print("文本内容:")
        for text in extracted_content["text_content"]:
            print(text)
        print("\nOCR处理图片内容:")
        for image_text in extracted_content["image_content"]:
            print(image_text)
        print("\n表格内容:")
        for table in extracted_content["table_content"]:
            for row in table:
                print(row)
            print()

    except ValueError as e:
        print(f"发生值错误：{str(e)}")

    except Exception as e:
        print(f"发生异常：{str(e)}")
