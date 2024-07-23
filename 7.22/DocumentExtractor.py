import os
import tempfile
from urllib.parse import urlparse
import requests
import base64
import json
import logging
from io import BytesIO
from docx import Document
from PIL import Image
from docx.text.paragraph import Paragraph
from docx.parts.image import ImagePart
from docx.oxml.shape import CT_Picture

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


class DocumentExtractor:
    """加载文档并进行OCR处理"""

    def __init__(self, file_path: str, paddleocr_url: str = "http://10.12.3.94:8956/ocr"):
        """初始化DocumentExtractor，包含文件路径和可选OCR URL"""
        self.doc = Document(file_path)  # 初始化文档对象
        self.file_path = file_path
        self.paddleocr_url = paddleocr_url

        if "~" in self.file_path:  # 如果文件路径中包含"~"，展开为绝对路径
            self.file_path = os.path.expanduser(self.file_path)

        # 如果文件是一个网络路径，下载到临时文件
        if not os.path.isfile(self.file_path) and self._is_valid_url(self.file_path):
            r = requests.get(self.file_path)

            if r.status_code != 200:
                raise ValueError(
                    f"Check the url of your file; returned status code {r.status_code}"
                )

            self.web_path = self.file_path
            self.temp_file = tempfile.NamedTemporaryFile(delete=False)
            self.temp_file.write(r.content)
            self.temp_file.close()
            self.file_path = self.temp_file.name

        elif not os.path.isfile(self.file_path):
            raise ValueError(
                f"File path {self.file_path} is not a valid file or url"
            )

    def __del__(self):
        """清理临时文件"""
        if hasattr(self, "temp_file"):
            os.remove(self.temp_file.name)

    def _is_valid_url(self, url: str) -> bool:
        """检查URL是否有效"""
        parsed = urlparse(url)
        return bool(parsed.netloc) and bool(parsed.scheme)

    def _pilimg_to_text(self, pil_img, paddleocr_url: str) -> str:
        """使用PaddleOCR将PIL图像转换为OCR处理的文本。"""
        buffered = BytesIO()
        pil_img.save(buffered, format=pil_img.format)
        image_data_bytes = buffered.getvalue()
        img_base64 = base64.b64encode(image_data_bytes).decode('utf-8')
        return self._get_ocr_result(img_base64, paddleocr_url)

    def _get_ocr_result(self, img_base64: str, paddleocr_url: str) -> str:
        """发送图像base64数据到PaddleOCR API并检索OCR结果"""
        headers = {"Content-Type": "application/json"}
        data = {"img_base64": img_base64}
        try:
            response = requests.post(url=paddleocr_url, headers=headers, data=json.dumps(data))
            response.raise_for_status()
            result = response.json()["rec_texts"][0]
            return ",".join(result)
        except requests.exceptions.RequestException as e:
            logging.error("Error performing OCR: %s", str(e))
            return ""

    def get_table_indices_and_contents(self):
        """获取文档中所有表格的起始索引位置及其内容"""
        table_data = []
        # 遍历文档的所有块元素
        for i, block in enumerate(self.doc.element.body):
            if block.tag.endswith('tbl'):
                # 获取表格内容
                table = self.doc.tables[len(table_data)]  # 根据当前表格计数获取表格对象
                table_content = []
                for row in table.rows:
                    row_content = [cell.text.strip() for cell in row.cells]
                    table_content.append(row_content)
                table_data.append({"index": i, "content": table_content})
        return table_data

    def extract(self, doc, title, extract_image=False):   # 默认不提取图片信息
        """Extract text, OCR processed images"""
        require_ = []
        title_index = self.get_title_index(doc, title)
        for i in range(len(title_index)-1):
            text_content = []
            image_content = []
            require_content = []
            start = title_index[i]
            end = title_index[i+1]

            for para_index, para in enumerate(doc.paragraphs):
                if start < para_index < end:
                    # 提取文本
                    text_content.append(para.text)
                    paras = doc.paragraphs[para_index]
                    # TODO：确保提取的内容正确，目前OCR不能用了  提取图片
                    img = paras._element.xpath('.//pic:pic')
                    if len(img) >= 1 and extract_image:
                        img: CT_Picture = img[0]
                        embed = img.xpath('.//a:blip/@r:embed')[0]
                        related_part: ImagePart = doc.part.related_parts[embed]
                        image: Image = related_part.image
                        result = self._pilimg_to_text(Image.open(BytesIO(image.blob)), self.paddleocr_url)
                        if result:
                            image_content.append('\n图片信息：'+result)
                            result = None
            require_content.append(''.join(text_content))

            if extract_image:
                require_content.append(''.join(image_content))
            require_.append(require_content)  # 包含了图片信息和文本信息

        return require_

    def get_title_index(self, doc, title):
        title_index = []
        for para_index, para in enumerate(doc.paragraphs):
            for i in range(len(title)):
                # 记录每一个需求的起始和结束索引
                if title[i] == para.text:
                    title_index.append(para_index)
                    # continue  # 跳过标题段落，直接进入提取内容阶段
        title_index.append(len(doc.paragraphs))  # 将最后一个位置加上去 防止索引超出
        return title_index

    def process_directory(self, doc):
        level1_dict = {}

        for para in doc.paragraphs:
            style_name = para.style.name.lower()
            text = para.text

            if 'toc' in style_name:
                level = int(style_name.split()[-1])

                if '\t' in text:
                    title, page_num = text.split('\t')
                else:
                    title = text
                    page_num = ''

                if level == 1:
                    level1_dict[title] = {'page_num': page_num, 'sub_dirs': {}}
                elif level == 2:
                    current_level1 = next(reversed(level1_dict))
                    level1_dict[current_level1]['sub_dirs'][title] = {'page_num': page_num, 'sub_dirs': {}}
                elif level == 3:
                    current_level1 = next(reversed(level1_dict))
                    current_level2 = next(reversed(level1_dict[current_level1]['sub_dirs']))
                    level1_dict[current_level1]['sub_dirs'][current_level2]['sub_dirs'][title] = {'page_num': page_num}

        return level1_dict

    def needed_directory_structure(self, directory_dict):
        subsub_text = []
        for level1_title, level1_info in directory_dict.items():
            if '需求' in level1_title:
                # print(f'Title: {level1_title}, Page Number: {level1_info["page_num"]}')
                for level2_title, level2_info in level1_info['sub_dirs'].items():
                    if len(level2_info['sub_dirs']) == 0:
                        #  print(f'  Sub Title: {level2_title}, Page Number: {level2_info["page_num"]}')
                        num, ttext = level2_title.split(' ')
                        subsub_text.append(ttext)
                    else:
                        for level3_title, level3_info in level2_info['sub_dirs'].items():
                            #print(f'    Sub Sub Title: {level3_title}, Page Number: {level3_info["page_num"]}')
                            num, ttext = level3_title.split(' ')
                            subsub_text.append(ttext)
        # print('需求的小标题信息', subsub_text)
        return subsub_text

    def text_table_image(self, title_index, table_info, require_):
        '''将提取到的文本、图片和表格信息按需求整合起来  √'''
        for i in range(len(title_index)-1):
            for j in table_info:
                if title_index[i] < j["index"] < title_index[i+1]:
                    # print("对应的需求提取的表格信息", title_index[i], j['index'], j['content'])
                    require_[i].append(j["content"])
        return require_

    def extract_document(self):
        '''
        传入上传的文档参数，提取目录信息，提取最低的一层子标题，
        提取具体的需求的文本、图片和表格等信息
        最终整合。
        '''
        doc = self.doc
        directory_dict = self.process_directory(doc)  # 提取目录
        sub_text = self.needed_directory_structure(directory_dict)  # 提取需要的需求小标题
        title_index = self.get_title_index(doc, sub_text)  # 提取需求小标题位置
        # print("需求在文档中出现的位置索引：", title_index)
        extracted_content = self.extract(doc, sub_text)
        # print(len(extracted_content))
        # for i in range(len(extracted_content)):
        #     print(extracted_content[i])
        table_info = self.get_table_indices_and_contents()  # 获取表格信息
        # print("表格数量", len(table_info))
        last_info = self.text_table_image(title_index, table_info, extracted_content)  # 将表格信息和其他信息整合
        # 已提取好
        # print('最终提取出来的需求数量：', len(last_info))
        # 返回需求小标题信息和提取到的需求信息传递给LLM
        return sub_text, last_info


if __name__ == '__main__':
    extractor = DocumentExtractor('D:\TestCases\航空云监控管理_需求规格说明书20240223.docx')
    sub_text, last_info = extractor.extract_document()
    # print("提取到的小标题:", sub_text)
    # print("提取到的内容:", last_info[1])
