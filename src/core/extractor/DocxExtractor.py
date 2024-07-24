import re
from docx import Document


class DocxExtractor:

    def extract_headings(self, document):

        headings = []

        # 提取标题和级别
        for i, paragraph in enumerate(document.paragraphs):
            if paragraph.style.name.startswith('Heading'):
                level = int(re.search(r'\d+', paragraph.style.name).group())
                text = paragraph.text.strip()
                headings.append((text, level, i))

        return headings

    def build_tree(self, headings):
        tree = []
        stack = []
        counters = {}

        for text, level, idx in headings:
            while stack and stack[-1][1] >= level:
                stack.pop()

            if level not in counters:
                counters[level] = 0
            if level == 1:
                counters[level] += 1
                for key in list(counters.keys()):
                    if key > level:
                        counters[key] = 0
            else:
                counters[level] += 1

            number = '.'.join(str(counters[l]) for l in range(1, level + 1) if l in counters)
            parent = stack[-1][0] if stack else None
            parent_number = stack[-1][3] if stack else None
            tree.append((text, parent, idx, number, parent_number))
            stack.append((text, level, idx, number))

        return tree

    def get_indices_and_content(self, document, tree):
        all_paragraphs = [p.text for p in document.paragraphs]
        indexed_tree = []

        for i, (title, parent, start_idx, number, parent_number) in enumerate(tree):
            end_idx = len(document.paragraphs) - 1
            for j in range(start_idx + 1, len(document.paragraphs)):
                if document.paragraphs[j].style.name.startswith('Heading'):
                    next_level = int(re.search(r'\d+', document.paragraphs[j].style.name).group())
                    if next_level <= int(re.search(r'\d+', document.paragraphs[start_idx].style.name).group()):
                        end_idx = j - 1
                        break
            content = ' '.join(all_paragraphs[start_idx + 1:end_idx + 1])
            indexed_tree.append((title, parent, start_idx, end_idx, content, number, parent_number))

        return indexed_tree

    def test(self):

        # 测试文档路径
        doc_path = r'D:\tmp\MyTestCaseGeneration\AI-test-case-generation\航空云监控管理_需求规格说明书20240223.docx'

        # 读取文档并提取标题
        headings = self.extract_headings(Document(doc_path))

        # 构建标题树
        tree = self.build_tree(headings)

        # 获取标题在文档中的索引位置并提取正文内容
        indexed_tree = self.get_indices_and_content(Document(doc_path), tree)

        # 输出结果
        for item in indexed_tree:
            print(f"Title: {item[0]}")
            print(f"Parent: {item[1]}")
            print(f"Start Index: {item[2]}")
            print(f"End Index: {item[3]}")
            print(f"Content: {item[4]}")
            print(f"Number: {item[5]}")
            print(f"Parent Number: {item[6]}")
            print("\n" + "-" * 50 + "\n")


docxExtractor = DocxExtractor()
