import os
import tempfile
from urllib.parse import urlparse
import requests
import base64
import json
import logging
from io import BytesIO
from docx import Document
from PIL import Image
from docx.parts.image import ImagePart
from docx.oxml.shape import CT_Picture
import httpx
from xinference.client import RESTfulClient

from core.extractor.DocxExtractor import docxExtractor
from core.mapper.DemandInfoMapper import DemandInfo
from core.service.DemandInfoService import demandInfoService
from core.util.SnowflakeUtil import snowflakeUtil

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


class DocumentExtractor:
    """加载文档并进行OCR处理"""

    def __init__(self, file_path: str, paddleocr_url: str = "http://10.12.3.94:8956/ocr"):
        """初始化DocumentExtractor，包含文件路径和可选OCR URL"""
        """初始化DocumentProcessor，包含本地化模型的客户端和模型实例"""
        self.doc = Document(file_path)
        self.file_path = file_path
        self.paddleocr_url = paddleocr_url

        # self.client = RESTfulClient("http://10.12.3.31:9997")
        # self.model = self.client.get_model("Qwen2-72B-Instruct-GPTQ-Int4")

        if "~" in self.file_path:  # 如果文件路径中包含"~"，展开为绝对路径
            self.file_path = os.path.expanduser(self.file_path)

        if not os.path.isfile(self.file_path) and self._is_valid_url(self.file_path):
            '''如果文件是一个网络路径，下载到临时文件'''
            r = requests.get(self.file_path)
            if r.status_code != 200:
                raise ValueError(f"Check the url of your file; returned status code {r.status_code}")

            self.web_path = self.file_path
            self.temp_file = tempfile.NamedTemporaryFile(delete=False)
            self.temp_file.write(r.content)
            self.temp_file.close()
            self.file_path = self.temp_file.name

        elif not os.path.isfile(self.file_path):
            raise ValueError(f"File path {self.file_path} is not a valid file or url")

    def __del__(self):
        """清理临时文件"""
        if hasattr(self, "temp_file"):
            os.remove(self.temp_file.name)

    def _is_valid_url(self, url: str) -> bool:
        """检查URL是否有效"""
        parsed = urlparse(url)
        return bool(parsed.netloc) and bool(parsed.scheme)

    def _pilimg_to_text(self, pil_img, paddleocr_url: str) -> str:
        """使用PaddleOCR将PIL图像转换为OCR处理的文本"""
        buffered = BytesIO()
        pil_img.save(buffered, format=pil_img.format)
        image_data_bytes = buffered.getvalue()
        img_base64 = base64.b64encode(image_data_bytes).decode('utf-8')
        return self._get_ocr_result(img_base64, paddleocr_url)

    def _get_ocr_result(self, img_base64: str, paddleocr_url: str) -> str:
        """发送图像base64数据到PaddleOCR API并检索OCR结果"""
        headers = {"Content-Type": "application/json"}
        data = {"img_base64": img_base64}
        try:
            response = requests.post(url=paddleocr_url, headers=headers, data=json.dumps(data))
            response.raise_for_status()
            result = response.json()["rec_texts"][0]
            return ",".join(result)
        except requests.exceptions.RequestException as e:
            logging.error("Error performing OCR: %s", str(e))
            return ""

    def process_directory(self, doc):
        '''处理目录信息'''
        level1_dict = {}

        for para in doc.paragraphs:
            style_name = para.style.name.lower()
            text = para.text

            if 'toc' in style_name:
                level = int(style_name.split()[-1])

                if '\t' in text:
                    title, page_num = text.split('\t')
                else:
                    title = text
                    page_num = ''

                if level == 1:
                    level1_dict[title] = {'page_num': page_num, 'sub_dirs': {}}
                elif level == 2:
                    current_level1 = next(reversed(level1_dict))
                    level1_dict[current_level1]['sub_dirs'][title] = {'page_num': page_num, 'sub_dirs': {}}
                elif level == 3:
                    current_level1 = next(reversed(level1_dict))
                    current_level2 = next(reversed(level1_dict[current_level1]['sub_dirs']))
                    level1_dict[current_level1]['sub_dirs'][current_level2]['sub_dirs'][title] = {'page_num': page_num}

        return level1_dict

    def needed_directory_structure(self, directory_dict):
        '''提取需求部分的目录信息'''
        subsub_text = []
        for level1_title, level1_info in directory_dict.items():
            if '需求' in level1_title:
                for level2_title, level2_info in level1_info['sub_dirs'].items():
                    if len(level2_info['sub_dirs']) == 0:
                        num, ttext = level2_title.split(' ')
                        subsub_text.append(ttext)
                    else:
                        for level3_title, level3_info in level2_info['sub_dirs'].items():
                            num, ttext = level3_title.split(' ')
                            subsub_text.append(ttext)
        return subsub_text

    def get_title_index(self, doc, title):
        '''获取需求小标题所在的具体文档位置'''
        title_index = []
        for para_index, para in enumerate(doc.paragraphs):
            for i in range(len(title)):
                if title[i] == para.text:
                    title_index.append(para_index)
        title_index.append(len(doc.paragraphs))
        # print(title_index)
        return title_index

    def get_table_indices_and_contents(self):
        """获取文档中所有表格的起始索引位置及其内容"""
        table_data = []
        for i, block in enumerate(self.doc.element.body):
            if block.tag.endswith('tbl'):
                table = self.doc.tables[len(table_data)]
                table_content = []
                for row in table.rows:
                    row_content = [cell.text.strip() for cell in row.cells]
                    table_content.append(row_content)
                table_data.append({"index": i, "content": table_content})
        # print(table_data)
        return table_data

    def extract_text_and_image(self, doc, title, extract_image='否'):
        """提取文本和OCR处理的图像"""
        required = []
        title_index = self.get_title_index(doc, title)
        for i in range(len(title_index) - 1):
            text_content = []
            image_content = []
            require_content = []
            start = title_index[i]
            end = title_index[i + 1]

            for para_index, para in enumerate(doc.paragraphs):
                if start < para_index < end:
                    text_content.append(para.text)  # 提取文本
                    paras = doc.paragraphs[para_index]
                    img = paras._element.xpath('.//pic:pic')  # 提取图片
                    if len(img) >= 1 and extract_image == '是':
                        img: CT_Picture = img[0]
                        embed = img.xpath('.//a:blip/@r:embed')[0]
                        related_part: ImagePart = doc.part.related_parts[embed]
                        image: Image = related_part.image
                        result = self._pilimg_to_text(Image.open(BytesIO(image.blob)), self.paddleocr_url)
                        if result:
                            image_content.append('\n图片信息：' + result)
                            result = None
            require_content.append(''.join(text_content))

            if extract_image == '是':
                require_content.append(''.join(image_content))
            required.append(require_content)

        return title_index, required

    def combineWithTable(self, title_index, table_info, require_content):
        """将提取到的文本、图片和表格信息按需求整合起来"""
        for i in range(len(title_index) - 1):
            for j in table_info:
                if title_index[i] < j["index"] < title_index[i + 1]:
                    require_content[i].append(j["content"])
        return require_content

    def extract_document(self, extract_image_flag='否'):
        """调用以上的函数，实现整份文档的内容提取"""
        doc = self.doc  # 初始化文档对象

        # 提取文档中的标题
        headings = docxExtractor.extract_headings(doc)
        # 构建标题树
        tree = docxExtractor.build_tree(headings)
        # 获取标题在文档中的索引位置并提取正文内容
        indexed_tree = docxExtractor.get_indices_and_content(self.doc, tree)

        # 批次ID
        batch_id = str(snowflakeUtil.get_id())
        # 遍历indexed_tree并写入表中
        for item in indexed_tree:
            # item 转化为DemandInfo实例
            demand = DemandInfo(
                id=str(snowflakeUtil.get_id()),
                batch_id=batch_id,
                title=item[0],
                parent=item[1],
                start_index=item[2],
                end_index=item[3],
                content=item[4],
                number=item[5],
                parent_number=item[6]
            )
            # 写入表中（这里下个版本再优化批量写入）
            demandInfoService.insert(demand)

        # 读取表数据并构建出参title, extract_content
        demandInfoList = demandInfoService.select_by_batch_id(batch_id)
        # 初始化出参
        title = []
        extract_content = []
        # 遍历插入出参
        for demandInfo in demandInfoList:
            title.append(demandInfo.title)
            extract_content.append(demandInfo.content)

        # doc = self.doc  # 初始化文档对象
        # directory_dict = self.process_directory(doc)  # 获取目录字典
        # sub_text = self.needed_directory_structure(directory_dict)  # 获取有关需求部分的目录
        # title_index, extracted_text_image = self.extract_text_and_image(doc, sub_text, extract_image)  # 获取文档中的文本内容和图片信息以及需求标题在文中的索引
        # table_info = self.get_table_indices_and_contents()  # 获取文档中的表格所有内容以及在文档中的开始索引
        # extract_content = self.combineWithTable(title_index, table_info, extracted_text_image)

        # 返回出参
        # print(title)
        # print(extract_content)
        return title, extract_content

    def join_md_content(self, title, extract_content):
        # 合并两个列表为Markdown表格格式
        markdown_table = "| 需求名称 | 需求说明 |\n| --- | --- |\n"
        for feature, detail in zip(title, extract_content):
            markdown_table += f"| {feature} | {detail} |\n"

        # 打印Markdown表格
        # print(markdown_table)
        return markdown_table

    async def qwen_extract_requirements(self, title, extract_content):
        """调用Qwen模型提取需求信息"""
        json_data = {
            "model": "Qwen/Qwen2-72B-Instruct",
            "max_tokens": 15000,
            "temperature": 0.5,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 1,
            "stop": [
                "<|im_end|>",
                "<|im_start|>"
            ],
            "messages": [
                {
                    "content": f"提取以下需求文档内容中的需求信息："
                               f"需求名称列表：{title}\n"
                               f"需求名称列表中各个需求对应的文本、图像、和表格信息{extract_content}\n"
                               f"输出格式：表格，分为需求名称、需求说明，总共两列。",
                    "role": "user"
                }
            ],
            "repetitive_penalty": 1
        }
        headers = {'Authorization': 'Bearer c9bd37061c2a92ea4d524aab6fd94d36aa108701f95811d3d2bbb8e41d0ff8d9'}
        async with httpx.AsyncClient() as client:
            response = await client.post("https://api.together.xyz/v1/chat/completions", json=json_data,
                                         headers=headers, timeout=300)
            response.raise_for_status()
            requirement_info = response.json()["choices"][0]["message"]["content"].strip()
            return requirement_info


if __name__ == '__main__':
    print('ok')
    # extractor = DocumentExtractor('D:\\TestCases\\航空云监控管理_需求规格说明书20240223.docx')
    # sub_text, last_info = extractor.extract_document()
